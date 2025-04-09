import os
import shutil
from fastapi import FastAPI, File, UploadFile, HTTPException, Request
from fastapi.responses import JSONResponse, FileResponse
import aiofiles
from pydub import AudioSegment
from google import genai
from dotenv import load_dotenv
from fastapi.responses import FileResponse
from docx import Document

load_dotenv()
# FastAPI Backend
app = FastAPI()

# Set the upload directory
UPLOAD_DIRECTORY = os.environ.get("Upload_dir", "uploaded_files")


@app.post("/uploadfile/")
async def upload_file(file: UploadFile = File(...)):
    if not os.path.exists(UPLOAD_DIRECTORY):
        os.makedirs(UPLOAD_DIRECTORY, exist_ok=True)
    """Handles file uploads and returns duration."""
    if not file:
        raise HTTPException(status_code=400, detail="No file uploaded")

    file_location = os.path.join(UPLOAD_DIRECTORY, file.filename)
    try:
        # Save the uploaded file
        async with aiofiles.open(file_location, "wb") as f:
            content = await file.read()
            await f.write(content)

        # Analyze the audio file
        audio = AudioSegment.from_file(file_location)
        duration_seconds = len(audio) / 1000  # Convert milliseconds to seconds
        hours = int(duration_seconds // 3600)
        minutes = int((duration_seconds % 3600) // 60)
        seconds = int(duration_seconds % 60)

        return JSONResponse(
            content={
                "filename": file.filename,
                "message": "File uploaded successfully",
                "audio_length": f"{hours}h {minutes}m {seconds}s",
                "audio_length_seconds": int(duration_seconds),
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File upload failed: {str(e)}")

@app.post("/process_audio/")
async def process_audio(request: Request):
    """Processes the audio file to trim it and optionally transcribe it."""
    try:
        data = await request.json()
        filename = data.get("filename")
        start_time = data.get("start_time")
        end_time = data.get("end_time")
        state = data.get("state", "trimming") 

        if not filename or start_time is None or end_time is None:
            raise HTTPException(
                status_code=400, detail="Missing required fields in the request payload"
            )

        # Construct the file path
        file_path = os.path.join(UPLOAD_DIRECTORY, filename)
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found")

        audio = AudioSegment.from_file(file_path)

        start_ms = int(start_time * 1000)
        end_ms = int(end_time * 1000)
        trimmed_audio = audio[start_ms:end_ms]

        processed_file_path = os.path.join(UPLOAD_DIRECTORY, f"processed_{filename}")
        trimmed_audio.export(processed_file_path, format="mp3")

        transcription_result = await generateTranscript(processed_file_path)

        doc = Document()
        doc.add_heading("Transcription", level=1)

        # Process the transcription result
        for line in transcription_result.splitlines():
            if line.startswith("M:"):  # Moderator's speech
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line)
                run.bold = True
            elif line.startswith("R:"):  # Responder's speech
                doc.add_paragraph(line)

        # Save the Word document
        os.makedirs("output_directory", exist_ok=True)
        transcription_file_path = os.path.join(
            "output_directory", f"transcription_{filename}.docx"
        )
        doc.save(transcription_file_path)

        # Return the Word document as a response
        return FileResponse(
            transcription_file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"transcription_{filename}.docx",
        )

    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Audio processing failed: {str(e)}"
        )


async def generateTranscript(processed_file_path):
    """Generates a transcript using GenAI, utilizing the Files API."""
    api_key = os.environ.get("GENAI_API_KEY")
    if not api_key:
        print("Error: GENAI_API_KEY environment variable not set.")
        return None

    client = genai.Client(api_key=api_key)

    try:
        uploaded_file = client.files.upload(file=processed_file_path)
        prompt = """As a highly skilled and experienced expert transcriber and translator (using Google's advanced language models), you will process an audio clip where speakers may use a mixture of English and other Indian languages (like Hindi or Telugu). Your primary goal is to produce a perfect English transcript. This means all spoken words, regardless of the original language, must be accurately translated into English. No other language should appear in the output.

                    Please follow these guidelines meticulously and check multiple times if you are doing it correctly:

                    * **Transcription and Translation:** Transcribe the audio into English, ensuring all non-English words are translated accurately and naturally.
                    * **Speaker Labeling:** Correctly label the speakers using the following format:
                        * M: for Moderator
                        * R: for Responder
                        * Example:
                            * M: Welcome to the meeting. Let's get started.
                            * R: Thank you. I have a few questions about the agenda.
                    * **Consecutive Turns:** When the same speaker has multiple turns consecutively, combine their dialogue into a single block of text under their respective speaker label.
                    * Example (this can be any similar pattern like this by the way):
                        * Input:
                            * M: Welcome to the meeting.
                            * M: Let's get started.
                            * R: Thank you.
                            * R: I have a few questions about the agenda.
                            * R: Can we discuss the timeline?
                        * Output:
                            * M: Welcome to the meeting. Let's get started.
                            * R: Thank you. I have a few questions about the agenda. Can we discuss the timeline?        
                    * **Apostrophes:** Use apostrophes correctly without any escape characters.
                    * **Language and Script:** The output must use only English script, and all words must be in English.
                    * **Grammar and Readability:** Ensure the transcript is grammatically correct and easy to read.

                    Your expertise in language processing and attention to detail are crucial for this task. Make sure, you go through every line carefully when ypu transcribe , the output should be in perfect english.
                    """

        result = client.models.generate_content(
            model="gemini-2.0-flash-exp",
            contents=[uploaded_file, prompt],
        )

        return result.text

    except Exception as e:
        print(f"An error occurred during transcript generation: {e}")
        return None
